#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md2docx_gbt.py — SOE Compliance Toolkit 统一排版生成器
Markdown → Word (.docx)，支持两种固定格式：

用法:
    python3 md2docx_gbt.py input.md output.docx [--type gongwen|court]

--type gongwen（默认）: GB/T 9704-2012 公文格式
    A4 上3.7/下3.5/左2.8/右2.6cm，行距固定28磅
    标题方正小标宋二号居中；正文仿宋三号首行缩进2字符
    一级节标题黑体三号；二级节标题楷体三号

--type court: 最高人民法院诉讼文书样式（通用版）
    A4 上2.54/下2.54/左3.17/右3.17cm，行距固定25磅
    标题宋体二号居中不加粗；正文宋体四号首行缩进2字符

支持标记:
    # 标题            → 按样式标题（居中）
    ## 一级节标题      → 样式 h1
    ### 二级节标题     → 样式 h2
    |表格             → 表格（首行表头加粗，正文按样式）
    > 说明/引用       → 正文段落
    >> 右对齐内容     → 右对齐（落款、日期用）
    - 列表项          → 正文段落
    ---               → 分隔线（忽略）
    普通段落           → 正文段落（首行缩进2字符）
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

# ---------- 字体常量 ----------
F_TITLE = 'FZXiaoBiaoSong-B05S'   # 方正小标宋
F_HEI = 'SimHei'                   # 黑体
F_KAI = 'KaiTi_GB2312'             # 楷体
F_FANG = 'FangSong_GB2312'         # 仿宋
F_SONG = 'SimSun'                  # 宋体

# ---------- 样式定义（内置默认值，format-spec.md 可覆盖） ----------
STYLES = {
    # 公文：GB/T 9704-2012
    'gongwen': {
        'title_font': F_TITLE, 'title_size': 22, 'title_bold': False,
        'h1_font': F_HEI, 'h2_font': F_KAI,
        'body_font': F_FANG, 'body_size': 16, 'body_line': 28,
        'table_size': 12,
        'margin_top': 3.7, 'margin_bottom': 3.5, 'margin_left': 2.8, 'margin_right': 2.6,
    },
    # 法院文书：最高法诉讼文书样式（通用）
    'court': {
        'title_font': F_SONG, 'title_size': 22, 'title_bold': False,
        'h1_font': F_SONG, 'h2_font': F_SONG,
        'body_font': F_SONG, 'body_size': 14, 'body_line': 25,
        'table_size': 10.5,
        'margin_top': 2.54, 'margin_bottom': 2.54, 'margin_left': 3.17, 'margin_right': 3.17,
    },
}
DEFAULT_STYLE = 'gongwen'

# 字段单位换算
PT_FIELDS = {'title_size', 'body_size', 'table_size', 'body_line'}
BOOL_FIELDS = {'title_bold'}
MARGIN_FIELDS = {'margin_top', 'margin_bottom', 'margin_left', 'margin_right'}


def load_format_spec(path=None):
    """从 format-spec.md 读取样式覆盖。返回 dict[style_name] -> {param: value}。"""
    if path is None:
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'format-spec.md')
    if not os.path.exists(path):
        return {}
    overrides = {}
    cur = None
    with open(path, encoding='utf-8') as f:
        for raw in f:
            line = raw.strip()
            if line.startswith('## '):
                cur = line[3:].strip()
                overrides.setdefault(cur, {})
                continue
            # 表格行 | key | value |
            if cur and line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.strip('|').split('|')]
                if len(cells) >= 2 and cells[0] and cells[0] != '参数':
                    key, val = cells[0], cells[1]
                    overrides[cur][key] = val
            # key: value 行
            elif cur and ':' in line and not line.startswith(('>', '-', '#')):
                key, _, val = line.partition(':')
                overrides[cur][key.strip()] = val.strip()
    return overrides


def resolve_style(style_name, overrides):
    """合并内置默认与外部覆盖，返回最终样式配置。"""
    base = dict(STYLES.get(style_name, STYLES[DEFAULT_STYLE]))
    ov = overrides.get(style_name, {})
    for k, v in ov.items():
        if k not in base:
            continue
        if k in BOOL_FIELDS:
            base[k] = str(v).lower() in ('true', '1', 'yes')
        elif k in PT_FIELDS:
            base[k] = float(v)
        elif k in MARGIN_FIELDS:
            base[k] = float(v)
        else:
            base[k] = v
    return base


def apply_style(doc, st):
    """按样式配置设置文档（页边距、默认字体）。"""
    sec = doc.sections[0]
    sec.top_margin = Cm(st['margin_top'])
    sec.bottom_margin = Cm(st['margin_bottom'])
    sec.left_margin = Cm(st['margin_left'])
    sec.right_margin = Cm(st['margin_right'])
    normal = doc.styles['Normal']
    normal.font.name = st['body_font']
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), st['body_font'])
    normal.font.size = Pt(st['body_size'])
    return st


def st_line(st):
    """行距（磅）转 Pt。"""
    return Pt(st['body_line'])


def set_font(run, name, size, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, st, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True,
             font=None, size=None, bold=False, line=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line if line is not None else st_line(st)
    if first_line:
        p.paragraph_format.first_line_indent = Pt((size or st['body_size']) * 2)
    run = p.add_run(text)
    set_font(run, font or st['body_font'], size or st['body_size'], bold)
    return p


def add_table(doc, rows, st):
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci in range(n_cols):
            txt = row[ci] if ci < len(row) else ''
            cell = cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            if ri == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p.add_run(txt), F_HEI, st['table_size'], bold=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
                set_font(p.add_run(txt), st['body_font'], st['table_size'])
            p.paragraph_format.line_spacing = Pt(22)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def is_table_sep(line):
    return bool(re.match(r'^\s*\|[\s:\-|]+\|\s*$', line)) and '-' in line


def parse_table(lines, i):
    rows = []
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            if is_table_sep(line):
                i += 1
                continue
            rows.append([c.strip() for c in line.strip('|').split('|')])
            i += 1
        else:
            break
    return rows, i


def convert(md_path, docx_path, style_name=DEFAULT_STYLE):
    overrides = load_format_spec()
    st = resolve_style(style_name, overrides)
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    apply_style(doc, st)

    i = 0
    in_code = False
    md_dir = os.path.dirname(os.path.abspath(md_path))
    while i < len(lines):
        line = lines[i].rstrip()
        s = line.strip()

        if not s or s == '---':
            i += 1
            continue

        # 图片 ![alt](path) —— 居中插入，宽度限制在版心内
        if s.startswith('![') and '](' in s and s.endswith(')'):
            alt = s[2:s.index('](')]
            path = s[s.index('](') + 2:-1]
            full = path if os.path.isabs(path) else os.path.join(md_dir, path)
            if os.path.exists(full):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run()
                run.add_picture(full, width=Cm(14.5))
                add_para(doc, alt, st, WD_ALIGN_PARAGRAPH.CENTER, False,
                         size=10, font=F_SONG)
            else:
                add_para(doc, f'[图片缺失: {alt}（{path}）]', st,
                         WD_ALIGN_PARAGRAPH.CENTER, False, bold=True)
            i += 1
            continue

        # 代码块（``` 包裹，等宽字体、左对齐、不缩进）
        if s.startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            add_para(doc, line, st, WD_ALIGN_PARAGRAPH.LEFT, False,
                     font='Consolas', size=max(st['body_size'] - 4, 9))
            i += 1
            continue

        if s.startswith('|'):
            rows, i = parse_table(lines, i)
            if rows:
                add_table(doc, rows, st)
            continue

        if s.startswith('# '):
            add_para(doc, s[2:].strip(), st, WD_ALIGN_PARAGRAPH.CENTER, False,
                     st['title_font'], st['title_size'], st['title_bold'])
            i += 1
            continue
        if s.startswith('## '):
            add_para(doc, s[3:].strip(), st, first_line=False, font=st['h1_font'])
            i += 1
            continue
        if s.startswith('### '):
            add_para(doc, s[4:].strip(), st, first_line=False, font=st['h2_font'])
            i += 1
            continue

        if s.startswith('>> '):
            add_para(doc, s[3:].strip(), st, WD_ALIGN_PARAGRAPH.RIGHT, False)
            i += 1
            continue
        if s.startswith('> '):
            add_para(doc, s[2:].strip(), st)
            i += 1
            continue
        if s.startswith('- '):
            add_para(doc, s[2:].strip(), st)
            i += 1
            continue

        add_para(doc, s, st)
        i += 1

    doc.save(docx_path)
    print(f'✅ 生成: {docx_path} (样式: {style_name})')


if __name__ == '__main__':
    args = sys.argv[1:]
    style = DEFAULT_STYLE
    if '--type' in args:
        idx = args.index('--type')
        style = args[idx + 1]
        del args[idx:idx + 2]
    if len(args) < 2:
        print(__doc__)
        sys.exit(1)
    convert(args[0], args[1], style)
